"""Document structure extraction service.

Extracts hierarchical structure from legal documents using PyMuPDF (primary)
or Azure Document Intelligence (optional, if configured). Detects sections,
headings, paragraphs, and generates structural citations for precise legal
grounding.

Optional environment variables (for Azure upgrade):
- AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT
- AZURE_DOCUMENT_INTELLIGENCE_KEY
"""

import asyncio
import hashlib
import io
import os
import re
import uuid
from dataclasses import dataclass, field
from typing import Any

# Timeout for Azure Document Intelligence API calls (seconds)
AZURE_TIMEOUT_SECONDS = 120


@dataclass
class SectionData:
    """Represents a section in the document hierarchy."""

    id: str
    parent_id: str | None
    section_number: str | None
    title: str | None
    level: int
    sequence: int
    path: list[str]
    start_page: int | None
    end_page: int | None
    content: str


@dataclass
class ChunkData:
    """Represents a chunk with citation data."""

    id: str
    section_id: str | None
    parent_chunk_id: str | None
    chunk_level: str  # 'section' or 'paragraph'
    chunk_index: int
    content: str
    content_hash: str
    citation: dict[str, Any]


@dataclass
class ExtractionResult:
    """Result of document structure extraction."""

    document_id: str
    sections: list[SectionData] = field(default_factory=list)
    chunks: list[ChunkData] = field(default_factory=list)
    normalized_markdown: str = ""
    extraction_quality: float = 0.0
    page_count: int = 0
    error: str | None = None


class StructureExtractor:
    """Extracts hierarchical structure from documents.

    Uses PyMuPDF for fast, reliable extraction from PDF and DOCX files.
    Falls back to Azure Document Intelligence if configured and requested.
    """

    # Regex patterns for section numbering
    SECTION_PATTERNS = [
        r"^(\d+\.)+\s*",  # 1.2.3 style
        r"^(\d+)\s+",  # 1 style
        r"^§\s*\d+",  # § 512 style
        r"^[IVXLCDM]+\.\s*",  # Roman numerals (I. II. III.)
        r"^[A-Z]\.\s*",  # Letter sections (A. B. C.)
        r"^Part\s+[IVXLCDM]+",  # Part I, Part II
        r"^Article\s+\d+",  # Article 1, Article 2
        r"^Section\s+\d+",  # Section 1, Section 2
        r"^Clause\s+\d+",  # Clause 1, Clause 2
    ]

    def __init__(self, use_hi_res: bool = True):
        """Initialize the structure extractor.

        Args:
            use_hi_res: Whether to prefer Azure Document Intelligence when available.
        """
        self.use_hi_res = use_hi_res

    async def extract(
        self,
        document_id: str,
        file_content: bytes,
        file_type: str,
    ) -> ExtractionResult:
        """Extract structure from a document.

        Uses Azure Document Intelligence if configured, otherwise falls back
        to PyMuPDF for PDFs and python-docx/mammoth for DOCX.

        Args:
            document_id: UUID of the document.
            file_content: Raw bytes of the document.
            file_type: File extension (pdf, docx, doc).

        Returns:
            ExtractionResult with sections, chunks, and normalized markdown.
        """
        # Validate file type
        supported_types = ("pdf", "docx", "doc")
        if file_type.lower() not in supported_types:
            return ExtractionResult(
                document_id=document_id,
                error=f"Unsupported file type: {file_type}. Supported: {', '.join(supported_types)}",
            )

        # Try Azure if configured and hi-res requested
        azure_endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
        azure_key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_KEY")

        if self.use_hi_res and azure_endpoint and azure_key:
            try:
                result = await self._extract_with_azure(
                    document_id, file_content, file_type, azure_endpoint, azure_key
                )
                if not result.error:
                    return result
                print(
                    f"[StructureExtractor] Azure failed, falling back to PyMuPDF: {result.error}"
                )
            except Exception as e:
                print(
                    f"[StructureExtractor] Azure exception, falling back to PyMuPDF: {e}"
                )

        # Primary extraction: PyMuPDF for PDFs, python-docx for DOCX
        if file_type.lower() == "pdf":
            return await self._extract_with_pymupdf(document_id, file_content)
        else:
            return await self._extract_with_docx(document_id, file_content, file_type)

    # ── PyMuPDF extraction (primary for PDFs) ────────────────────────────

    async def _extract_with_pymupdf(
        self,
        document_id: str,
        file_content: bytes,
    ) -> ExtractionResult:
        """Extract structure from a PDF using PyMuPDF.

        Uses font size analysis to detect headings and builds a section hierarchy.
        """

        def _do_extract():
            import fitz  # PyMuPDF

            doc = fitz.open(stream=file_content, filetype="pdf")

            # First pass: collect all text blocks with font info to determine heading thresholds
            all_blocks: list[dict[str, Any]] = []
            font_sizes: list[float] = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)[
                    "blocks"
                ]

                for block in blocks:
                    if block.get("type") != 0:  # text blocks only
                        continue

                    for line in block.get("lines", []):
                        text_parts = []
                        line_sizes: list[float] = []

                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                text_parts.append(text)
                                line_sizes.append(span.get("size", 12.0))

                        full_text = " ".join(text_parts).strip()
                        if not full_text:
                            continue

                        avg_size = (
                            sum(line_sizes) / len(line_sizes) if line_sizes else 12.0
                        )
                        is_bold = any(
                            "bold" in span.get("font", "").lower()
                            for span in line.get("spans", [])
                        )

                        all_blocks.append(
                            {
                                "text": full_text,
                                "page": page_num + 1,
                                "size": avg_size,
                                "bold": is_bold,
                            }
                        )
                        font_sizes.append(avg_size)

            if not all_blocks:
                return ExtractionResult(
                    document_id=document_id,
                    normalized_markdown="",
                    extraction_quality=0.0,
                    page_count=len(doc),
                )

            # Determine heading thresholds from font size distribution
            font_sizes.sort()
            if len(font_sizes) > 0:
                # Body text is the most common font size
                from collections import Counter

                size_counts = Counter(round(s, 1) for s in font_sizes)
                body_size = size_counts.most_common(1)[0][0]
            else:
                body_size = 12.0

            # Build sections and chunks
            sections: list[SectionData] = []
            chunks: list[ChunkData] = []
            markdown_lines: list[str] = []
            section_stack: list[SectionData] = []
            chunk_index = 0
            paragraph_index = 0

            for block in all_blocks:
                text = block["text"]
                page = block["page"]
                size = block["size"]
                bold = block["bold"]

                # Detect if this is a heading
                is_heading = False
                level = 1

                # Significantly larger than body = heading
                size_ratio = size / body_size if body_size > 0 else 1.0

                if size_ratio >= 1.5:
                    is_heading = True
                    level = 1
                elif size_ratio >= 1.25:
                    is_heading = True
                    level = 2
                elif size_ratio >= 1.1 and bold:
                    is_heading = True
                    level = 2
                elif bold and len(text) < 120:
                    # Bold short lines are likely headings
                    is_heading = True
                    level = 3

                # Section number patterns override font-based detection
                if self._extract_section_number(text) and len(text) < 200:
                    is_heading = True
                    level = self._detect_heading_level(None, text)

                if is_heading:
                    # Pop sections at same or higher level
                    while section_stack and section_stack[-1].level >= level:
                        section_stack.pop()

                    parent_id = section_stack[-1].id if section_stack else None
                    path = [s.title or "" for s in section_stack] + [text]

                    section = SectionData(
                        id=str(uuid.uuid4()),
                        parent_id=parent_id,
                        section_number=self._extract_section_number(text),
                        title=text,
                        level=level,
                        sequence=len(sections),
                        path=path,
                        start_page=page,
                        end_page=page,
                        content="",
                    )
                    sections.append(section)
                    section_stack.append(section)

                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=section.id,
                            parent_chunk_id=None,
                            chunk_level="section",
                            chunk_index=chunk_index,
                            content=text,
                            content_hash=self._generate_content_hash(text),
                            citation={
                                "page": page,
                                "section_path": path,
                                "paragraph_index": None,
                                "heading": text,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1

                    markdown_lines.append(f"{'#' * level} {text}")
                    markdown_lines.append("")
                else:
                    # Regular paragraph
                    current_section = section_stack[-1] if section_stack else None
                    section_path = [s.title or "" for s in section_stack]

                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=current_section.id if current_section else None,
                            parent_chunk_id=None,
                            chunk_level="paragraph",
                            chunk_index=chunk_index,
                            content=text,
                            content_hash=self._generate_content_hash(text),
                            citation={
                                "page": page,
                                "section_path": section_path,
                                "paragraph_index": paragraph_index,
                                "heading": current_section.title
                                if current_section
                                else None,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1
                    paragraph_index += 1

                    markdown_lines.append(text)
                    markdown_lines.append("")

            normalized_markdown = "\n".join(markdown_lines).strip()
            page_count = len(doc)

            # Quality: sections detected = higher quality
            quality = 0.85 if sections else 0.65

            print(
                f"[StructureExtractor] PyMuPDF extracted {len(sections)} sections, "
                f"{len(chunks)} chunks from {page_count} pages"
            )

            return ExtractionResult(
                document_id=document_id,
                sections=sections,
                chunks=chunks,
                normalized_markdown=normalized_markdown,
                extraction_quality=quality,
                page_count=page_count,
                error=None,
            )

        try:
            return await asyncio.to_thread(_do_extract)
        except Exception as e:
            return ExtractionResult(
                document_id=document_id,
                error=f"PyMuPDF extraction failed: {e}",
            )

    # ── DOCX extraction ──────────────────────────────────────────────────

    async def _extract_with_docx(
        self,
        document_id: str,
        file_content: bytes,
        file_type: str,
    ) -> ExtractionResult:
        """Extract structure from a DOCX/DOC file using python-docx."""

        def _do_extract():
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_content))

            sections: list[SectionData] = []
            chunks: list[ChunkData] = []
            markdown_lines: list[str] = []
            section_stack: list[SectionData] = []
            chunk_index = 0
            paragraph_index = 0

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower() if para.style else ""

                # Detect headings from style
                is_heading = False
                level = 1

                if "heading" in style_name:
                    is_heading = True
                    # Extract level from style name (e.g., "Heading 2")
                    match = re.search(r"(\d+)", style_name)
                    level = int(match.group(1)) if match else 1
                elif "title" in style_name:
                    is_heading = True
                    level = 1
                elif self._extract_section_number(text) and len(text) < 200:
                    # Detect by section number pattern
                    is_heading = True
                    level = self._detect_heading_level(None, text)

                if is_heading:
                    while section_stack and section_stack[-1].level >= level:
                        section_stack.pop()

                    parent_id = section_stack[-1].id if section_stack else None
                    path = [s.title or "" for s in section_stack] + [text]

                    section = SectionData(
                        id=str(uuid.uuid4()),
                        parent_id=parent_id,
                        section_number=self._extract_section_number(text),
                        title=text,
                        level=level,
                        sequence=len(sections),
                        path=path,
                        start_page=None,
                        end_page=None,
                        content="",
                    )
                    sections.append(section)
                    section_stack.append(section)

                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=section.id,
                            parent_chunk_id=None,
                            chunk_level="section",
                            chunk_index=chunk_index,
                            content=text,
                            content_hash=self._generate_content_hash(text),
                            citation={
                                "page": None,
                                "section_path": path,
                                "paragraph_index": None,
                                "heading": text,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1

                    markdown_lines.append(f"{'#' * level} {text}")
                    markdown_lines.append("")
                else:
                    current_section = section_stack[-1] if section_stack else None
                    section_path = [s.title or "" for s in section_stack]

                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=current_section.id if current_section else None,
                            parent_chunk_id=None,
                            chunk_level="paragraph",
                            chunk_index=chunk_index,
                            content=text,
                            content_hash=self._generate_content_hash(text),
                            citation={
                                "page": None,
                                "section_path": section_path,
                                "paragraph_index": paragraph_index,
                                "heading": current_section.title
                                if current_section
                                else None,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1
                    paragraph_index += 1

                    markdown_lines.append(text)
                    markdown_lines.append("")

            # Handle tables
            for table in doc.tables:
                table_md = self._docx_table_to_markdown(table)
                if table_md:
                    markdown_lines.append(table_md)
                    markdown_lines.append("")

            normalized_markdown = "\n".join(markdown_lines).strip()
            quality = 0.85 if sections else 0.65

            print(
                f"[StructureExtractor] DOCX extracted {len(sections)} sections, "
                f"{len(chunks)} chunks"
            )

            return ExtractionResult(
                document_id=document_id,
                sections=sections,
                chunks=chunks,
                normalized_markdown=normalized_markdown,
                extraction_quality=quality,
                page_count=0,
                error=None,
            )

        try:
            return await asyncio.to_thread(_do_extract)
        except Exception as e:
            return ExtractionResult(
                document_id=document_id,
                error=f"DOCX extraction failed: {e}",
            )

    # ── Azure Document Intelligence (optional) ───────────────────────────

    async def _extract_with_azure(
        self,
        document_id: str,
        file_content: bytes,
        file_type: str,
        azure_endpoint: str,
        azure_key: str,
    ) -> ExtractionResult:
        """Extract using Azure Document Intelligence.

        Args:
            document_id: UUID of the document.
            file_content: Raw bytes of the document.
            file_type: File extension.
            azure_endpoint: Azure Document Intelligence endpoint URL.
            azure_key: Azure Document Intelligence API key.

        Returns:
            ExtractionResult with sections, chunks, and normalized markdown.
        """
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
        from azure.core.credentials import AzureKeyCredential

        print(f"[StructureExtractor] Using Azure Document Intelligence for {file_type}")

        client = DocumentIntelligenceClient(
            endpoint=azure_endpoint,
            credential=AzureKeyCredential(azure_key),
        )

        # Analyze document with prebuilt-layout model (best for structure)
        # Run in thread to avoid blocking the event loop
        def _analyze():
            poller = client.begin_analyze_document(
                "prebuilt-layout",
                AnalyzeDocumentRequest(bytes_source=file_content),
            )
            return poller.result()

        try:
            result = await asyncio.wait_for(
                asyncio.to_thread(_analyze),
                timeout=AZURE_TIMEOUT_SECONDS,
            )
        except asyncio.TimeoutError:
            return ExtractionResult(
                document_id=document_id,
                error=f"Azure Document Intelligence timed out after {AZURE_TIMEOUT_SECONDS} seconds",
            )

        # Extract sections from paragraphs with roles
        sections: list[SectionData] = []
        chunks: list[ChunkData] = []
        markdown_lines: list[str] = []

        section_stack: list[SectionData] = []
        chunk_index = 0
        paragraph_index = 0

        # Process paragraphs
        if result.paragraphs:
            for para in result.paragraphs:
                role = para.role if hasattr(para, "role") else None
                content = para.content.strip() if para.content else ""

                if not content:
                    continue

                # Get page number
                page = None
                if para.bounding_regions and len(para.bounding_regions) > 0:
                    page = para.bounding_regions[0].page_number

                # Check if this is a heading/title
                if role in ("title", "sectionHeading", "heading"):
                    # Determine level
                    level = 1
                    if role == "sectionHeading":
                        level = self._detect_heading_level(None, content)

                    # Pop sections at same or higher level
                    while section_stack and section_stack[-1].level >= level:
                        section_stack.pop()

                    parent_id = section_stack[-1].id if section_stack else None
                    path = [s.title or "" for s in section_stack] + [content]

                    section = SectionData(
                        id=str(uuid.uuid4()),
                        parent_id=parent_id,
                        section_number=self._extract_section_number(content),
                        title=content,
                        level=level,
                        sequence=len(sections),
                        path=path,
                        start_page=page,
                        end_page=page,
                        content="",
                    )
                    sections.append(section)
                    section_stack.append(section)

                    # Add section chunk
                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=section.id,
                            parent_chunk_id=None,
                            chunk_level="section",
                            chunk_index=chunk_index,
                            content=content,
                            content_hash=self._generate_content_hash(content),
                            citation={
                                "page": page,
                                "section_path": path,
                                "paragraph_index": None,
                                "heading": content,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1

                    markdown_lines.append(f"{'#' * level} {content}")
                    markdown_lines.append("")
                else:
                    # Regular paragraph
                    current_section = section_stack[-1] if section_stack else None
                    section_path = [s.title or "" for s in section_stack]

                    chunks.append(
                        ChunkData(
                            id=str(uuid.uuid4()),
                            section_id=current_section.id if current_section else None,
                            parent_chunk_id=None,
                            chunk_level="paragraph",
                            chunk_index=chunk_index,
                            content=content,
                            content_hash=self._generate_content_hash(content),
                            citation={
                                "page": page,
                                "section_path": section_path,
                                "paragraph_index": paragraph_index,
                                "heading": current_section.title
                                if current_section
                                else None,
                                "context_before": None,
                                "context_after": None,
                            },
                        )
                    )
                    chunk_index += 1
                    paragraph_index += 1

                    markdown_lines.append(content)
                    markdown_lines.append("")

        # Process tables
        if result.tables:
            for table in result.tables:
                table_md = self._azure_table_to_markdown(table)
                markdown_lines.append(table_md)
                markdown_lines.append("")

        normalized_markdown = "\n".join(markdown_lines).strip()
        page_count = result.pages[-1].page_number if result.pages else 0

        # Calculate quality (Azure generally produces high-quality results)
        quality = 0.9 if sections else 0.7

        print(
            f"[StructureExtractor] Azure extracted {len(sections)} sections, {len(chunks)} chunks"
        )

        return ExtractionResult(
            document_id=document_id,
            sections=sections,
            chunks=chunks,
            normalized_markdown=normalized_markdown,
            extraction_quality=quality,
            page_count=page_count,
            error=None,
        )

    # ── Helpers ───────────────────────────────────────────────────────────

    def _azure_table_to_markdown(self, table: Any) -> str:
        """Convert Azure table to markdown format."""
        if not table.cells:
            return ""

        max_row = max(cell.row_index for cell in table.cells) + 1
        max_col = max(cell.column_index for cell in table.cells) + 1

        grid = [["" for _ in range(max_col)] for _ in range(max_row)]

        for cell in table.cells:
            content = cell.content.strip() if cell.content else ""
            grid[cell.row_index][cell.column_index] = content

        lines = []
        for i, row in enumerate(grid):
            line = "| " + " | ".join(row) + " |"
            lines.append(line)
            if i == 0:
                lines.append("| " + " | ".join(["---"] * max_col) + " |")

        return "\n".join(lines)

    def _docx_table_to_markdown(self, table: Any) -> str:
        """Convert a python-docx table to markdown format."""
        rows = table.rows
        if not rows:
            return ""

        lines = []
        for i, row in enumerate(rows):
            cells = [cell.text.strip() for cell in row.cells]
            line = "| " + " | ".join(cells) + " |"
            lines.append(line)
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        return "\n".join(lines)

    def _detect_heading_level(self, element: Any, title_text: str) -> int:
        """Detect heading level from content patterns.

        Args:
            element: Element (unused, kept for API compatibility).
            title_text: The heading text.

        Returns:
            Heading level (1-6).
        """
        if re.match(r"^\d+\.\d+\.\d+", title_text):
            return 3
        elif re.match(r"^\d+\.\d+", title_text):
            return 2
        elif re.match(r"^\d+\s", title_text):
            return 1
        elif re.match(r"^[A-Z]\.", title_text):
            return 2
        elif re.match(r"^[ivxIVX]+\.", title_text):
            return 3

        return 1

    def _extract_section_number(self, title_text: str) -> str | None:
        """Extract section number from heading text."""
        for pattern in self.SECTION_PATTERNS:
            match = re.match(pattern, title_text)
            if match:
                return match.group(0).strip()
        return None

    def _generate_content_hash(self, content: str) -> str:
        """Generate SHA-256 hash for content verification."""
        return hashlib.sha256(content.encode()).hexdigest()


def generate_content_hash(content: str) -> str:
    """Generate SHA-256 hash for content verification."""
    return hashlib.sha256(content.encode()).hexdigest()


def verify_content_hash(content: str, stored_hash: str) -> bool:
    """Verify content matches stored hash."""
    return generate_content_hash(content) == stored_hash
