"""Production-grade document text extraction service.

Extracts text from PDF, DOCX, and TXT files using:
- PDF: Hybrid extraction (native text + DeepSeek OCR for scanned pages)
- DOCX: mammoth for HTML conversion (preserves structure) with python-docx fallback
- TXT: Simple decoding with charset detection
"""

import asyncio
import io
import os
from pathlib import Path
from typing import Any


class DocumentProcessor:
    """Production-grade document text extraction.

    Features:
    - Hybrid PDF extraction: native text + OCR fallback for scanned pages
    - DOCX extraction with structure preservation
    - Intelligent chunking for embedding
    """

    SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt", "rtf"}

    # Target chunk size in characters (approximately 500 tokens)
    DEFAULT_CHUNK_SIZE = 2000
    # Overlap between chunks for context continuity
    DEFAULT_CHUNK_OVERLAP = 200

    # Minimum text per PDF page to skip OCR
    MIN_TEXT_THRESHOLD = 50

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        deepseek_api_key: str | None = None,
        deepseek_base_url: str | None = None,
        use_ocr: bool = True,
    ):
        """Initialize the document processor.

        Args:
            chunk_size: Target size for text chunks in characters.
            chunk_overlap: Overlap between chunks in characters.
            deepseek_api_key: API key for DeepSeek OCR (for scanned PDFs).
            deepseek_base_url: DeepSeek API base URL.
            use_ocr: Whether to use OCR for scanned PDF pages.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.deepseek_api_key = deepseek_api_key or os.getenv("DEEPSEEK_API_KEY")
        self.deepseek_base_url = deepseek_base_url or os.getenv("DEEPSEEK_BASE_URL")
        self.use_ocr = use_ocr and bool(self.deepseek_api_key)
        self._ocr_client: Any = None

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from document content (sync wrapper).

        Args:
            file_content: Raw bytes of the document file.
            file_type: File type (pdf, docx, txt).

        Returns:
            Extracted text content.

        Raises:
            ValueError: If file type is not supported.
        """
        return asyncio.run(self.extract_text_async(file_content, file_type))

    async def extract_text_async(self, file_content: bytes, file_type: str) -> str:
        """Extract text from document content (async).

        Args:
            file_content: Raw bytes of the document file.
            file_type: File type (pdf, docx, txt).

        Returns:
            Extracted text content.

        Raises:
            ValueError: If file type is not supported.
        """
        file_type = file_type.lower().strip(".")

        # Normalize doc to docx
        if file_type == "doc":
            file_type = "docx"

        if file_type not in self.SUPPORTED_TYPES:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: {', '.join(self.SUPPORTED_TYPES)}"
            )

        if file_type == "pdf":
            return await self._extract_pdf_async(file_content)
        elif file_type in ("docx", "doc"):
            return self._extract_docx(file_content)
        elif file_type == "rtf":
            return self._extract_rtf(file_content)
        else:  # txt
            return self._extract_txt(file_content)

    def extract_text_from_file(self, file_path: str | Path) -> str:
        """Extract text from a file path.

        Args:
            file_path: Path to the document file.

        Returns:
            Extracted text content.
        """
        path = Path(file_path)
        file_type = path.suffix.lower().strip(".")
        with open(path, "rb") as f:
            return self.extract_text(f.read(), file_type)

    async def _extract_pdf_async(self, content: bytes) -> str:
        """Extract text from PDF with OCR fallback for scanned pages.

        Uses a hybrid approach:
        1. Try native text extraction first (fast)
        2. Fall back to DeepSeek Vision OCR for pages with little/no text

        Args:
            content: Raw PDF bytes.

        Returns:
            Extracted text.
        """
        # Try PyMuPDF first (better extraction), fall back to pypdf
        import importlib.util

        if importlib.util.find_spec("fitz") is not None:
            return await self._extract_pdf_pymupdf(content)
        else:
            # Fall back to pypdf (no OCR capability)
            return self._extract_pdf_pypdf(content)

    async def _extract_pdf_pymupdf(self, content: bytes) -> str:
        """Extract text using PyMuPDF with OCR fallback."""
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        pages_text: list[str] = []
        ocr_needed_pages: list[tuple[int, Any]] = []  # (page_num, pixmap)

        # First pass: Native text extraction
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text().strip()

            if len(text) >= self.MIN_TEXT_THRESHOLD:
                pages_text.append(text)
            else:
                # Mark for OCR, store pixmap
                if self.use_ocr:
                    mat = fitz.Matrix(2.0, 2.0)  # 144 DPI
                    pix = page.get_pixmap(matrix=mat)
                    ocr_needed_pages.append((page_num, pix.tobytes("png")))
                    pages_text.append(None)  # Placeholder
                else:
                    # No OCR available, use what we have
                    pages_text.append(text if text else "[No text extracted]")

        doc.close()

        # Second pass: OCR for scanned pages
        if ocr_needed_pages and self.use_ocr:
            await self._ocr_pages(ocr_needed_pages, pages_text)

        # Filter and join
        return "\n\n".join(text for text in pages_text if text)

    async def _ocr_pages(
        self,
        pages: list[tuple[int, bytes]],
        pages_text: list[str | None],
    ) -> None:
        """Run OCR on pages that need it.

        Args:
            pages: List of (page_num, image_bytes) tuples.
            pages_text: List of page texts to update in place.
        """
        from .deepseek_ocr import DeepSeekOCR

        async with DeepSeekOCR(
            api_key=self.deepseek_api_key,
            base_url=self.deepseek_base_url,
        ) as ocr:
            for page_num, image_data in pages:
                try:
                    ocr_text = await ocr.extract_text_from_image(image_data, "png")
                    if ocr_text and ocr_text.strip() not in (
                        "[BLANK PAGE]",
                        "[UNREADABLE]",
                    ):
                        pages_text[page_num] = ocr_text
                    else:
                        pages_text[page_num] = ""
                except Exception as e:
                    print(f"Warning: OCR failed for page {page_num + 1}: {e}")
                    pages_text[page_num] = ""

    def _extract_pdf_pypdf(self, content: bytes) -> str:
        """Extract text using pypdf (fallback, no OCR)."""
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX with structure preservation.

        Tries mammoth first (better HTML conversion preserving structure),
        falls back to python-docx.

        Args:
            content: Raw DOCX bytes.

        Returns:
            Extracted text.
        """
        # Try mammoth first - better structure preservation
        try:
            return self._extract_docx_mammoth(content)
        except ImportError:
            pass

        # Fall back to python-docx
        return self._extract_docx_python_docx(content)

    def _extract_docx_mammoth(self, content: bytes) -> str:
        """Extract text from DOCX using mammoth.

        Mammoth converts DOCX to HTML/markdown, preserving:
        - Headings and their hierarchy
        - Lists (ordered and unordered)
        - Tables
        - Bold/italic (optional)
        """
        import mammoth
        from markdownify import markdownify

        # Convert DOCX to HTML
        result = mammoth.convert_to_html(io.BytesIO(content))
        html = result.value

        # Convert HTML to markdown for clean text with structure
        text = markdownify(html, heading_style="ATX", strip=["a", "img"])

        # Clean up excessive whitespace while preserving structure
        lines = []
        prev_blank = False
        for line in text.split("\n"):
            line = line.rstrip()
            is_blank = not line.strip()

            if is_blank:
                if not prev_blank:
                    lines.append("")
                prev_blank = True
            else:
                lines.append(line)
                prev_blank = False

        return "\n".join(lines).strip()

    def _extract_docx_python_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx (fallback)."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element == element and para.text.strip():
                        parts.append(para.text)
                        break
            # Handle tables
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        parts.append(self._extract_table(table))
                        break

        return "\n\n".join(parts)

    def _extract_table(self, table: Any) -> str:
        """Extract text from a DOCX table as markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            # Add header separator
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)

        return "\n".join(rows)

    def _extract_rtf(self, content: bytes) -> str:
        """Extract text from RTF content.

        Args:
            content: Raw RTF bytes.

        Returns:
            Extracted text.
        """
        try:
            from striprtf.striprtf import rtf_to_text

            # Decode RTF
            rtf_text = content.decode("utf-8", errors="ignore")
            return rtf_to_text(rtf_text)
        except ImportError:
            # Basic RTF stripping fallback
            text = content.decode("utf-8", errors="ignore")
            # Remove RTF control words (basic)
            import re

            text = re.sub(r"\\[a-z]+\d*\s?", "", text)
            text = re.sub(r"[{}]", "", text)
            return text.strip()

    def _extract_txt(self, content: bytes) -> str:
        """Extract text from TXT content with charset detection.

        Args:
            content: Raw text bytes.

        Returns:
            Decoded text.
        """
        # Try common encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue

        # Last resort: ignore errors
        return content.decode("utf-8", errors="ignore")

    def chunk_text(self, text: str) -> list[str]:
        """Split text into overlapping chunks for embedding.

        Attempts to split on paragraph/section boundaries when possible.

        Args:
            text: Full document text.

        Returns:
            List of text chunks.
        """
        if not text or not text.strip():
            return []

        # If text is smaller than chunk size, return as single chunk
        if len(text) <= self.chunk_size:
            return [text.strip()]

        chunks: list[str] = []

        # Split on double newlines (paragraphs/sections)
        paragraphs = text.split("\n\n")
        current_chunk: list[str] = []
        current_length = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_length = len(para)

            # If single paragraph exceeds chunk size, split it
            if para_length > self.chunk_size:
                # Save current chunk if any
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_length = 0

                # Split long paragraph into smaller chunks
                chunks.extend(self._split_long_text(para))
                continue

            # Check if adding this paragraph would exceed chunk size
            if current_length + para_length + 2 > self.chunk_size:
                # Save current chunk
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))

                # Start new chunk with overlap
                if self.chunk_overlap > 0 and current_chunk:
                    # Take last portion of previous chunk for overlap
                    overlap_text = "\n\n".join(current_chunk)[-self.chunk_overlap :]
                    current_chunk = [overlap_text, para]
                    current_length = len(overlap_text) + para_length + 2
                else:
                    current_chunk = [para]
                    current_length = para_length
            else:
                current_chunk.append(para)
                current_length += para_length + 2

        # Don't forget the last chunk
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks

    def _split_long_text(self, text: str) -> list[str]:
        """Split a long text into chunks at sentence boundaries.

        Args:
            text: Long text to split.

        Returns:
            List of chunks.
        """
        chunks: list[str] = []
        current = ""

        # Split on sentence boundaries
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)

        for sentence in sentences:
            if len(current) + len(sentence) > self.chunk_size:
                if current:
                    chunks.append(current.strip())
                current = sentence
            else:
                current += " " + sentence if current else sentence

        if current:
            chunks.append(current.strip())

        return chunks


# Convenience function for quick extraction
async def extract_document_text(
    file_content: bytes,
    file_type: str,
    use_ocr: bool = True,
) -> str:
    """Extract text from a document.

    Args:
        file_content: Raw bytes of the document.
        file_type: File extension (pdf, docx, txt).
        use_ocr: Whether to use OCR for scanned PDFs.

    Returns:
        Extracted text.
    """
    processor = DocumentProcessor(use_ocr=use_ocr)
    return await processor.extract_text_async(file_content, file_type)
